from flask import Flask, request, render_template, send_file, redirect, url_for
from docx import Document
import os
import datetime
import zipfile
import shutil
import re
from concurrent.futures import ThreadPoolExecutor

app = Flask(__name__)

# Create uploads and outputs directory if they don't exist
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

@app.route('/')
def upload_file():
    # Pass the current date to render the HTML template
    current_date = datetime.datetime.now().strftime('%Y-%m-%d')
    return render_template('upload.html', current_date=current_date)

@app.route('/upload', methods=['POST'])
def process_file():
    # Get the uploaded file
    file = request.files['file']
    date_input = request.form.get('date')  # Get the date from the form

    # Validate the provided date or use the current date
    try:
        if date_input:
            formatted_date = datetime.datetime.strptime(date_input, '%Y-%m-%d').strftime('%d-%m-%Y')
            date_message = f"Using provided date: {formatted_date}"
        else:
            formatted_date = datetime.datetime.now().strftime('%d-%m-%Y')
            date_message = f"No date provided. Using current date: {formatted_date}"
    except ValueError:
        # If the date format is invalid, default to the current date
        formatted_date = datetime.datetime.now().strftime('%d-%m-%Y')
        date_message = f"Invalid date entered. Defaulting to current date: {formatted_date}"

    # Save the uploaded file to the uploads directory
    if file and file.filename.endswith('.docx'):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)

        # Process the uploaded DOCX file
        process_docx(file_path, formatted_date)

        # Create a zip of the generated .docx files
        zip_filename = create_zip_of_docs()

        # Serve the zip file to download
        response = send_file(zip_filename, as_attachment=True)

        # Clean up the uploaded file and output folder
        cleanup_uploads_and_outputs()

        # Add the date message to the response headers for user feedback
        response.headers["X-Date-Message"] = date_message

        return response

    return redirect(url_for('upload_file'))


# Function to replace Turkish characters with English equivalents
def convert_turkish_characters(text):
    turkish_chars = 'çÇğĞıİöÖşŞüÜ'
    english_chars = 'cCgGiIoOsSuU'
    translation_table = str.maketrans(turkish_chars, english_chars)
    return text.translate(translation_table)

def keep_after_last_vowel(s):
    vowels = "AEIİOÖUÜaeıioöuü"  # Define vowels (including Turkish vowels)
    last_vowel_pos = -1  # Initialize with -1 to handle cases with no vowels

    # Loop through the string to find the last vowel position
    for i, char in enumerate(s):
        if char in vowels:
            last_vowel_pos = i  # Update position whenever a vowel is found

    # If a vowel was found, slice from that position onward; else, return the original string
    return s[last_vowel_pos +2:] if last_vowel_pos != -1 else s

def process_row(row, formatted_date, output_folder, template_path, template_path2):
    if 'BİRLİK' not in str(row.cells[1].text).strip() and  str(row.cells[5].text).strip() != '':
        # Extract and process the name
        name = re.sub(r'\s+', ' ', str(row.cells[5].text)).strip()
        if '-' in name:
            name_parts = name.split(' ')
            sup = name_parts[-1].strip()  # Take the last element after splitting
            original_name = name.replace(sup, '').strip()  # Remove the last part from original name
            # Ensure `sup` contains a hyphen
            if '-' not in sup:
                sup = ''.join(name_parts[-2:])
                original_name = ' '.join(name_parts[:-2])

            if len(sup.split('-')[0].strip()) > 4:
                first_part = keep_after_last_vowel(sup.split('-')[0])
                if len(first_part) > 1:
                    sup = first_part + '-' + sup.split('-')[1]
                else :
                    sup = sup.split('-')[0][-2:] + '-' + sup.split('-')[1]
                original_name = name.replace(sup, '').strip()  # Remove the last part from original name
        else:
            original_name = name  # If no hyphen, keep the name as is
            sup = ''

        # Clean the 'doktor' variable of non-ASCII characters and illegal filesystem characters
        doktor = str(row.cells[16].text).strip()


        # Create a new document from the template within this thread
        if 'FAKO' in str(row.cells[12].text).strip() and 'VRC' not in str(row.cells[12].text).strip():
            new_doc = Document(template_path2)
        else:
            new_doc = Document(template_path)

        # Replace placeholders with data
        new_doc.tables[0].rows[0].cells[1].text = formatted_date
        new_doc.tables[0].rows[2].cells[0].text = sup
        new_doc.tables[0].rows[2].cells[1].text = original_name
        new_doc.tables[0].rows[2].cells[2].text = str(row.cells[11].text).strip()
        new_doc.tables[0].rows[2].cells[3].text = str(row.cells[12].text).strip()

        if doktor in ['KENAN SÖNMEZ', 'AYŞE GÜL KOÇAK ALTINTAŞ', 'AYŞE GÜL KOÇAK', 'AYŞEGÜL KOÇAK ALTINTAŞ', 'AYŞEGÜL KOÇAK', 'MEHMET YASİN TEKE', 'MEHMET ÇITIRIK','MUSTAFA İLKER TOKER', 'BERRAK ŞEKERYAPAN GEDİZ', 'BERRAK ŞEKERYAPAN', 'YASEMİN ÖZDAMAR EROL']:
            new_doc.tables[0].rows[5].cells[1].text = 'PROF.DR.' + doktor
        elif doktor in ['BURCU KAZANCI', 'FATMA ÇORAK','FATMA ÇORAK EROĞLU', 'NURETTİN BAYRAM', 'PINAR ÇİÇEK', 'EREN EKİCİ']:
            new_doc.tables[0].rows[5].cells[1].text = 'DOÇ.DR.' + doktor
        else: new_doc.tables[0].rows[5].cells[1].text = 'OP.DR.' + doktor

        # Take only the first two parts of the name
        doktor_parts = doktor.split()[:2]
        doktor_short = " ".join(doktor_parts)

        # Convert Turkish characters and replace spaces with underscores
        if '-' in name:
            folder_name = "ASST CASEs"
        else:
            folder_name = convert_turkish_characters(doktor_short)

        # Create the path for the doctor's subfolder
        doktor_subfolder = os.path.join(output_folder, folder_name)

        # Sequentially create the subfolder if it doesn't exist
        os.makedirs(doktor_subfolder, exist_ok=True)

        # Sequentially save the generated document to the outputs folder
        #file_name = convert_turkish_characters(original_name)
        new_doc.save(os.path.join(doktor_subfolder, f"{original_name}.docx"))

def process_docx(file_path, formatted_date):
    # Load the uploaded DOCX file
    document = Document(file_path)

    # Set output_folder to the main OUTPUT_FOLDER
    output_folder = app.config['OUTPUT_FOLDER']
    os.makedirs(output_folder, exist_ok=True)

    # Gather all rows (skip the first 3 header rows)
    rows = []
    for table in document.tables:
        rows.extend(table.rows[3:])

    # Use ThreadPoolExecutor to process rows in parallel (within safe boundaries)
    with ThreadPoolExecutor() as executor:
        executor.map(
            lambda row: process_row(row, formatted_date, output_folder, "template.docx", "template2.docx"),
            rows
        )

def create_zip_of_docs():
    output_folder = app.config['OUTPUT_FOLDER']

    # Create a zip file of the generated .docx files
    zip_filename = os.path.join(output_folder, f"hasta_kart_{datetime.datetime.now().strftime('%Y-%m-%d')}.zip")

    with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED, compresslevel=9) as docx_zip:
        for root, dirs, files in os.walk(output_folder):
            for file in files:
                if file.endswith(".docx") and file != ".docx":
                    # Write the file to the zip archive, preserving the subfolder structure
                    arcname = os.path.relpath(os.path.join(root, file), output_folder)
                    docx_zip.write(os.path.join(root, file), arcname=arcname)

                    # After adding to the zip, delete the .docx file to free up space
                    os.remove(os.path.join(root, file))

    return zip_filename


### Call clean up user button
@app.route('/cleanup', methods=['POST'])
def cleanup():
    cleanup_uploads_and_outputs()
    return '', 204  # Return 'No Content' status to indicate success

def cleanup_uploads_and_outputs():
    # Define directories
    upload_folder = app.config['UPLOAD_FOLDER']
    output_folder = app.config['OUTPUT_FOLDER']

    # Remove all contents and the directory itself for uploads
    try:
        if os.path.exists(upload_folder):
            shutil.rmtree(upload_folder)
            print("Uploads folder and its contents removed.")
    except Exception as e:
        print(f"Error during uploads folder cleanup: {e}")


    # Remove all contents and the directory itself for outputs
    try:
        if os.path.exists(output_folder):
            shutil.rmtree(output_folder)
            print("Outputs folder contents removed.")

            if os.path.exists(output_folder):
                os.rmdir(output_folder)  # Attempt to remove it explicitly
                print("Outputs folder itself removed.")
    except Exception as e:
        print(f"Error during outputs folder cleanup: {e}")

    # Recreate empty folders for future requests
    os.makedirs(upload_folder, exist_ok=True)
    os.makedirs(output_folder, exist_ok=True)
    print("Uploads and Outputs folders recreated.")
if __name__ == "__main__":
    app.run(debug=True)